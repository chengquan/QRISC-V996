#!/usr/bin/env python3
# 生成公众号文章 Word 版(python-docx),内嵌 3 张架构图,中文字体。
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

HERE = os.path.dirname(os.path.abspath(__file__))
BLUE = RGBColor(0x1a,0x4d,0x7a); GREY = RGBColor(0x66,0x66,0x66)
CJK = "微软雅黑"   # Word 打开时若无此字体会回退,正文照常显示

doc = Document()

# 全局字体
st = doc.styles["Normal"]; st.font.size = Pt(11); st.font.name="Calibri"
st.element.rPr.rFonts.set(qn('w:eastAsia'), CJK)

def set_cjk(run):
    run.font.name="Calibri"; run._element.rPr.rFonts.set(qn('w:eastAsia'),CJK)

def H(text,size=16,color=BLUE,before=14,after=6,center=False):
    p=doc.add_paragraph(); p.space_before=Pt(before)
    if center: p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(text); r.bold=True; r.font.size=Pt(size); r.font.color.rgb=color; set_cjk(r)
    p.paragraph_format.space_before=Pt(before); p.paragraph_format.space_after=Pt(after)
    return p

def P(text,size=11,color=None,italic=False,center=False):
    p=doc.add_paragraph()
    if center: p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(text); r.font.size=Pt(size); r.italic=italic; set_cjk(r)
    if color: r.font.color.rgb=color
    p.paragraph_format.space_after=Pt(6); return p

def BULLET(text,size=11):
    p=doc.add_paragraph(style="List Bullet"); r=p.add_run(text); r.font.size=Pt(size); set_cjk(r); return p

def CODE(text):
    p=doc.add_paragraph(); p.paragraph_format.left_indent=Inches(0.2)
    r=p.add_run(text); r.font.name="Consolas"; r.font.size=Pt(9.5); r.font.color.rgb=RGBColor(0x22,0x33,0x44)
    sh=p._p.get_or_add_pPr(); from docx.oxml import OxmlElement
    shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),'F2F4F6'); sh.append(shd)
    p.paragraph_format.space_after=Pt(6); return p

def FIG(name,caption,width=6.0):
    doc.add_picture(os.path.join(HERE,name),width=Inches(width))
    doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
    c=doc.add_paragraph(); c.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=c.add_run(caption); r.font.size=Pt(9); r.italic=True; r.font.color.rgb=GREY; set_cjk(r)
    c.paragraph_format.space_after=Pt(10)

# ============ 标题 ============
t=doc.add_paragraph(); t.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=t.add_run("我把一整套 Linux,跑在了自己“造”的 RISC-V 芯片上"); r.bold=True
r.font.size=Pt(20); r.font.color.rgb=BLUE; set_cjk(r)
P("—— QRISC-V996:一个从核到操作系统全部自编的开源 RISC-V SoC 仿真平台",
  size=11,color=GREY,italic=True,center=True)

# ============ 引子 ============
H("先说结论",15)
P("不用买开发板、不用流片,只要一台装了 Linux 的电脑(Windows 的 WSL 也行),"
  "你就能看着一颗 RISC-V 处理器——从上电、跑引导程序、启动 Linux 内核、"
  "一直到出现 “~ #” 命令行提示符——全过程都在你眼前的“数字电路仿真”里发生。"
  "而且这颗“芯片”和这套 Linux,从硬件电路到操作系统,全部是从源代码自己编译出来的。")
P("这个项目叫 QRISC-V996。下面用大白话讲清楚:它是什么、为什么有意思、你能拿它干什么。")

# ============ 图1 全景 ============
H("一、它到底是什么?",15)
P("一句话:一个“能开机、能交互、能跑程序”的完整 RISC-V 电脑——只不过这台电脑活在仿真里。")
FIG("fig1_overview.png","图 1　全景:你敲命令 → 仿真里的 RISC-V 芯片 → 运行自编的 Linux")
P("拆开看是四层,每一层都是开源、可自己重建的:")
BULLET("处理器核:biRISC-V,一颗双发射的 RV32IMA 核(带指令/数据缓存、内存管理单元 MMU)——相当于芯片的“大脑”。")
BULLET("外设:UART 串口、定时器、GPIO、SPI、中断控制器——全是真正的硬件电路描述(RTL),不是软件假装的。")
BULLET("操作系统:一套从源码编译的 Linux 5.4 + BusyBox 文件系统。")
BULLET("引导程序:SBI 固件,负责把内核拉起来、并补上这颗核没有的几条指令。")
P("这四层用 Verilator(把电路“翻译”成可运行程序的工具)做成一个仿真器,"
  "我们叫它 tb_soc。你通过一个串口控制台 GUI 跟里面的 Linux 打字交互,就像连着一块真开发板。")

# ============ 二、为什么有意思 ============
H("二、为什么这件事有意思?",15)
P("因为它把通常“看不见、摸不着”的东西，全摊开在你面前了：")
BULLET("看得见硬件:一条指令怎么取、缓存怎么命中、外设寄存器怎么被读写,都能用波形看到每一个时钟周期。")
BULLET("全栈自编:从 Verilog 电路,到 GCC 交叉编译器,到 Linux 内核,到根文件系统——没有黑盒,每一层都能改、能重建。")
BULLET("真实但安全:它是真实的电路逻辑(不是“差不多得了”的行为模型),却跑在电脑里,改错了也不会烧板子。")
P("一句话:它是学习“计算机到底怎么从一堆门电路变成能跑 Linux 的系统”的绝佳活教材。")

# ============ 三、SoC 内部 ============
H("三、芯片内部长什么样?",15)
P("处理器核通过 AXI 总线,接到一个“互联”上;互联再按地址把请求分发给 DRAM 和各个外设。"
  "每个外设都有一个固定的“门牌号”(地址),软件往那个地址读写,就等于操作真实硬件。")
FIG("fig2_soc.png","图 2　SoC 内部:核 + AXI 互联 + 真 RTL 外设(每个外设一个地址)")
P("这里踩过两个很“硬核”的坑,也是这个项目最有价值的部分之一:")
BULLET("AXI ID 路由:核的取指和访存走两条总线,响应必须按编号送回正确的那条,否则取指会永久卡死、根本起不来。")
BULLET("中断电平:外部中断信号必须“电平跟随”,否则中断只来一次核就僵住——这是让真中断驱动的串口能工作的关键。")
P("这两个都是真实的电路 Bug,定位它们要靠看波形一拍一拍地追,修好后系统才能一路启动到命令行。")

# ============ 四、写程序 ============
H("四、你能拿它干什么?——写程序",15)
P("平台自带一个 SDK,两种玩法:")
P("① 裸机程序(没有操作系统):直接在核上跑,直接读写外设寄存器去点灯(GPIO)、发串口、收发 SPI。"
  "启动快、最适合验证硬件和看波形。一条命令就能编译并跑:")
CODE("./build_run.sh examples/gpio.c")
P("② Linux 程序(在系统里当进程跑):用标准 C 写,经 /dev/mem 操作外设。这里有个很实用的小发明——“虚拟磁盘”。")

# ============ 图3 虚拟磁盘 ============
FIG("fig3_vdisk.png","图 3　虚拟磁盘:改一个程序几秒搞定,不用重建内核")
P("通常你改一个 Linux 程序,要把它重新“烤”进内核镜像里,重建一次要一两分钟。"
  "虚拟磁盘的做法是:把程序单独打包,放到内存里内核不用的一块区域;开机时系统再用 mmap 把它取出来放进 /opt。"
  "于是改程序只要重打这块磁盘(几秒),内核完全不用动——迭代效率天差地别。")

# ============ 五、怎么玩 ============
H("五、上手有多快?",15)
P("装好依赖、编译一次仿真器(一两分钟),就能启动:")
CODE("python3 gui/biriscv_soc_console.py     # 点“启动”,等到 ~ # 就能敲命令了")
P("到了命令行,uname、ls、cat /proc/cpuinfo 都能跑;cpuinfo 里会看到 isa: rv32ima、mmu: sv32"
  "——这就是 biRISC-V 核的真实身份证。想看硬件波形,勾一下“录波形”,GPIO/SPI 引脚的每一次跳变都看得清清楚楚。")

# ============ 结尾 ============
H("写在最后",15)
P("QRISC-V996 把“一台 RISC-V 电脑”从最底层的电路,到最上层的 Linux 命令行,完整地、"
  "可复现地摆在了你面前。它不追求跑得多快(周期级仿真本就慢),而是追求“透明”——"
  "每一层都看得见、改得动、能从零重建。")
P("如果你想真正搞懂“计算机是怎么从晶体管一路变成能跑 Linux 的系统”,"
  "这是一个少见的、能让你亲手把整条链路走一遍的项目。",color=GREY,italic=True)
P("")
P("致谢:处理器核 biRISC-V、外设 riscv_soc、引导器 riscv-linux-boot 均来自开源社区 "
  "ultraembedded;操作系统基于 Linux 5.4 与 BusyBox。",size=9,color=GREY)

out=os.path.join(HERE,"QRISC-V996-公众号文章.docx")
doc.save(out)
print("✅ 生成",out,"(%d KB)"%(os.path.getsize(out)//1024))
